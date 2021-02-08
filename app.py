import requests
from bs4 import BeautifulSoup
from flask import Flask, render_template, request,send_file,session,redirect, url_for
from flask_mysqldb import MySQL
import sqlite3 as sql
import time
import atexit
from apscheduler.schedulers.background import BackgroundScheduler
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'MyDB'

mysql = MySQL(app)
app.secret_key = 'aniruddhabhattacharyasatulbhattsparshamit'


def RepresentsInt(s):
    try:
        int(s)
        return True
    except ValueError:
        return False

def call():
    from firebase import firebase
    firebase = firebase.FirebaseApplication('https://mausamkendradehradun.firebaseio.com/', None)
    data = {
        'Name': 'Pradip Bhattacharya',
        'Email': 'coolsamrat576@gmail.com',
        'Phone': 7983645854
    }
    result = firebase.post('mausamkendradehradun/videolibrary', data)

    #result = firebase.get('mausamkendradehradun/videolibrary', '')
    print(result)
    print("hi")


@app.route('/firebase')
def firbase():
    print("hi")
    call()
    return "Firebase"


def show_Forcast(district):
    data = requests.get("http://nwp.imd.gov.in/uttaranchal.txt")
    data_content = data.text
    District_data = data_content.split('DISTRICT :')

    district_Dictionary = {"UTTARKASHI": 1, "CHAMOLI": 2, "PITHORAGARH": 3, "BAGESHWAR": 4, "ALMORA": 5, "CHAMPAWAT": 6,
                           "UDHAMSINGH-NGR": 7, "NAINITAL": 8, "PAURI": 9, "HARIDWAR": 10, "RUDRAPRAYAG": 11,
                           "TEHRI-GARHWAL": 12, "DEHRADUN": 13}
    index = district_Dictionary[district]
    unique_Data = District_data[index]
    unique_Data = unique_Data.replace("-", "")
    unique_Data_List = unique_Data.split("\n")

    figures_Dict = {}
    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = [], [], [], [], [], [], [], []
    for i in range(2, len(unique_Data_List) - 3):
        figures_List = []
        unique_Data_List[i] = unique_Data_List[i].replace("     ", "")
        new_List = unique_Data_List[i].split(" ")
        figures_List.clear()

        for j in range(len(new_List)):
            if (RepresentsInt(new_List[j])):
                figures_List.append(new_List[j])

            if (i == 2):
                rainfall = figures_List
            if (i == 3):
                maxt = figures_List
            if (i == 4):
                mint = figures_List
            if (i == 5):
                tcc = figures_List
            if (i == 6):
                marh = figures_List
            if (i == 7):
                mirh = figures_List
            if (i == 8):
                ws = figures_List
            if (i == 9):
                wd = figures_List
                # print(figures_List)

    return rainfall, maxt, mint, tcc, marh, mirh, ws, wd

def scrapper(url):
    page = requests.get(url)

    soup = BeautifulSoup(page.content, 'html.parser')
    data = []

    table_body = soup.find_all('td')

    heading_1 = soup.find_all('b')
    heading_2 = soup.find_all('font', attrs={'color': 'blue'})
    heading_final = heading_1[0].text + heading_2[0].text + ", " + heading_1[1].text

    col2 = soup.find_all('td', attrs={'width': '60%'})
    col2_heading = heading_1[3].text
    rows_col2 = soup.find_all('font', attrs={'size': '1+'})
    ranged = 0
    listofdata = []
    for i in range(0, len(rows_col2), 2):

        row_data = rows_col2[i].text.strip() + " : " + rows_col2[i + 1].text.strip('\n')
        if rows_col2[i].text.strip() == "Todays Sunset (IST)":
            listofdata.append(rows_col2[i + 1].text.strip('\n'))
        if rows_col2[i].text.strip() == "Tommorows Sunrise (IST)":
            listofdata.append(rows_col2[i + 1].text.strip('\n'))
        if rows_col2[
            i].text.strip() == "24 Hours Rainfall (mm) (Recorded from 0830 hrs IST of yesterday to 0830 hrs IST of today)":
            listofdata.append(rows_col2[i + 1].text.strip('\n'))
        if "Moonrise" in row_data:
            ranged = i + 2
            break

    col3_heading = ""
    for i in range(ranged, ranged + 4):
        col3_heading = col3_heading + " | " + rows_col2[i].text.replace('\n', ' ').replace(' ', '')

    ranged = ranged + 4
    try:
        tableofdata = []
        for i in range(ranged, len(rows_col2), 4):
            tableofdata.append([rows_col2[i].text.replace('\n', ' ').replace(' ', ''),
                                rows_col2[i + 1].text.replace('\n', ' ').replace(' ', ''),
                                rows_col2[i + 2].text.replace('\n', ' ').replace(' ', ''),
                                rows_col2[i + 3].text.replace('\n', ' ').replace('  ', '')])
    except:
        pass
    return listofdata, tableofdata

def table(top):
    try:
        document = Document('/var/www/html/flaskproject/pdf/special_press_release.docx')
        table = document.tables[top]
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        #print(data)
        keys2 = []
        values2 = []

        for i in data:
            keys2.append(list(i.keys()))
            values2.append(list(i.values()))
        print(keys2)
        return keys2[0],values2
    except IndexError as e:
        a = []
        b = []
        return a, b

@app.route('/special')
def return_files_tut():
    #return send_file('/var/www/html/flaskproject/pdf/SPECIAL_PRESS_RELEASE.pdf', attachment_filename='SPECIAL_PRESS_RELEASE.pdf')
    document = Document('/var/www/html/flaskproject/pdf/special_press_release.docx')
    lst=[]
    for block in iter_block_items(document):
        lst.append(block.text.strip() if isinstance(block, Paragraph) else '@=-><-=@')
    keybhai,valuebhai = table(0)
    keybhai2, valuebhai2 = table(1)
    return render_template('agr.html', textdata=lst, keybhai=keybhai,valuebhai=valuebhai,keybhai2=keybhai2, valuebhai2=valuebhai2)
    #except Exception as e:
	#return str(e)
@app.route("/mansarovar")
def radar():
    yatraheading = "Mansarover Yatra"
    return render_template("agromet.html",yatraheading=yatraheading)
@app.route("/chardham")
def agromet():
    yatraheading = "Char Dham Yatra"
    return render_template("agromet.html",yatraheading=yatraheading)

def rainfall(city):
    import PyPDF2
    import re
    pdfFileObj = open(r'DISTRICT_FORECAST.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj, strict=False)
    number_of_pages = pdfReader.getNumPages()
    pageObj = pdfReader.getPage(0)
    page_content = pageObj.extractText()
    warning = ""
    found = False
    alert = ""

    page_content = page_content.replace("\n", "")

    page_content = page_content.split('District')[2]

    district = ["Uttarkashi", "Chamoli", "Rudraprayag", "Tehri", "Pauri", "Dehradun",
                "Haridwar", "Pithoragarh", "Bageshwar", "Almora", "Nainital", "Champawat", "U.S."]

    page_content = page_content.split(" ")

    dist_count = 0
    dist_dict = {}

    while (dist_count < len(district)):
        for i in range(len(page_content)):
            if district[dist_count] in page_content[i]:
                if district[dist_count] == "Tehri" or district[dist_count] == "Pauri":
                    dist_dict[district[dist_count]] = page_content[i + 2: i + 12]
                elif district[dist_count] == "U.S.":
                    district[dist_count] += " Nagar"
                    dist_dict[district[dist_count]] = page_content[i + 2: i + 12]
                else:
                    dist_dict[district[dist_count]] = page_content[i + 1: i + 11]
        dist_count += 1
    listj = []
    for i in dist_dict:
        if i == city:
            listj = ((dist_dict[i]))
            break
    return listj

@app.route('/ff')
def ff():
    citylistdata = ['Uttarkashi', 'Chamoli', 'Rudraprayag', 'Tehri', 'Pauri', 'Dehradun', 'Haridwar',
                    'Pithoragarh', 'Bageshwar', 'Almora', 'Nainital', 'Champawat', 'U.S. Nagar']
    citylistdata = sorted(citylistdata)
    listj = rainfall('Dehradun')
    return render_template('ff.html',citylistdata=citylistdata,listj=listj,city = 'Dehradun')

@app.route('/ffsearch',methods=['POST'])
def ffsearch():
    search_city = request.form['cityname']
    citylistdata = ['Uttarkashi', 'Chamoli', 'Rudraprayag', 'Tehri', 'Pauri', 'Dehradun', 'Haridwar',
                    'Pithoragarh', 'Bageshwar', 'Almora', 'Nainital', 'Champawat', 'U.S. Nagar']
    citylistdata = sorted(citylistdata)
    listj = rainfall(search_city)
    print(listj)
    return render_template('ff.html',citylistdata=citylistdata,listj=listj,city = search_city)


@app.route('/aniruddha')
def show_static_pdf():
    return send_file(r'WEATHER_FORECAST.pdf', attachment_filename='WEATHER_FORECAST.pdf')

@app.route('/sf')
def sf():
    from docx.api import Document

    document = Document('WEATHER_FORECAST.docx')
    table = document.tables[1]

    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

    dates = []
    hindiPrediction = []
    engPrediction = []

    for i in range(0, 5):
        prediction = list(data[i].values())
        dates.append(prediction[0])
        hindiPrediction.append(prediction[1])
        engPrediction.append(prediction[2])

    dates2 = []
    hindiWarning = []
    engWarning = []

    for i in range(8, 13):  # warning starts from 9th
        warning = list(data[i].values())
        dates2.append(warning[0])
        hindiWarning.append(warning[1])
        engWarning.append(warning[2])

    return render_template('sf.html',dates=dates,engPrediction=engPrediction,hindiPrediction=hindiPrediction,dates2=dates2,engWarning=engWarning,hindiWarning=hindiWarning)



def city_weather(city):
    city_Codes = {"Pithoragarh" :10007, "Pantnagar" :42148, "Nainital" :42146,
     "Hemkund Sahib" : 10057, "Champawat" :10010, "Mukteshwar" : 42147, "Almora" : 10008, "Chamoli" : 42117, "Joshimath" : 42116,
                  "Govindghat" : 10055, "Badrinath" :10047, "Devprayag" : 10052, "Srinagar" : 10062, "Rudraprayag" : 10060,
                  "Gauchar" : 10054, "Augustmuni" : 10046, "Ukhimat" : 10063, "Soneprayag" : 10061, "Kedarnath" : 0,
                  "Haridwar" : 10009, "Rishikesh" : 10059, "Chamba" : 10050, "Tehri" : 42114, "Dehradun" : 42111,
                  "Mussoorie" : 42112, "Nainbagh" : 10058, "Chinyalisaur" : 10051, "Lakhamandal" : 10065, "Barkot" : 10048,
                  "Uttarkashi" : 42108, "Bhatwari" : 10049, "Yamunotri" : 10064, "Harsil" : 10056, "Gangotri" : 10053
                 }
    cleardaytext = [
        "sunny day",
        "clear sky",
        "mainly clear sky",
        "mainly clear sky becoming partly cloudy towards a/e",
        "mainly clear sky becoming partly cloudy towards afternoon or evening",
        "mainly clear sky becoming partly cloudy towards e/n",
        "mainly clear sky becoming partly cloudy towards evening or night"

    ]
    cloudytext = [
        "partly cloudy sky",
        "partly cloudy sky towards afternoon/evening",
        "partly cloudy sky towards afternoon or evening",
        "partly cloudy sky with haze",
        "generally cloudy sky",
        "partly cloudy sky becoming generally cloudy towards afternoon or evening or night",
        "mainly clear sky becoming partly cloudy towards afternoon or evening or night",
        "generally cloudy towards afternoon or evening",
        "mainly clear sky becoming partly cloudy towards afternoon or evening",
        "generally cloudy sky with haze",
        "overcast sky",
        "partly cloudy sky with possibility of development of thunder/lightning",
        "partly cloudy sky with possibility of development of thunder or lightning",
        "partly cloudy sky with possibility of development of thunder lightning",
        "generally cloudy sky with possibility of development of thunder/lightning",
        "generally cloudy sky with possibility of development of thunder or lightning",
        "partly cloudy sky with possibility of rain or thunderstorm/dust storm",
        "partly cloudy sky with possibility of rain or thunderstorm or dust storm",
        "generally cloudy sky with possibility of rain",
        "generally cloudy sky with possibility of rain or thunderstorm",
        "generally cloudy sky with possibility of rain or dust storm",
        "partly cloudy sky with possibility of rain",
        "partly cloudy sky with possibility of rain or thunderstorm",
        "partly cloudy sky with possibility of rain or dust storm",

        "generally cloudy sky with possibility of rain or thunderstorm or dust storm",
    ]
    thunderstormtext = [
        "dust storm/thunderstorm",
        "partly cloudy sky with possibility of rain or thunderstorm/dust storm",
        "partly cloudy sky with possibility of rain or thunderstorm or dust storm",
        "partly cloudy sky with possibility of rain or thunderstorm or duststorm",
        "generally cloudy sky with possibility of rain or thunderstorm/dust storm",
        "generally cloudy sky with possibility of rain or thunderstorm or dust storm",
        "generally cloudy sky with possibility of rain or thunderstorm or duststorm",
        "thunderstorm",
        "partly cloudy sky with thundery development",
        "thunderstorm with rain",
        "thunderstorm with squall",
        "thunderstorm with hail",
        "thunderstorm with squall/hail",
        "thunderstorm with squall or hail",
        "dust storm/thunderstorm with squall/hail",
        "dust storm/thunderstorm with squall or hail",
    
    ]
    raintext = [
        "partly cloudy sky in the morning hours becomming generally cloudy sky towards evening/night with possibility of rain or thundershowers accompanied with squall",
        "partly cloudy sky in the morning hours becoming generally cloudy sky toward evening/night with possibility of rain or thundershowers accompanied with squall/hail",
        "partly cloudy sky in the morning hours becoming generally cloudy sky towards afternoon/evening with possibility of rain or thundershowers accompanied with squall/hail",
        "partly cloudy sky in the morning hours becoming generally cloudy sky towards afternoon/evening with possibility of rain or thundershowers accompanied with squall",
        "generally cloudy sky with light rain",
        "rain or thundershowers",
        "rain",
        "generally cloudy sky with light rain or drizzle",
        "rain or thundershowers would occur towards afternoon or evening",
        "rain or thundershowers would occur towards evening or night",
        "rain or thundershowers with strong gusty winds",
        "partly cloudy sky with one or two spells of rain or thundershowers",
        'generally cloudy sky with moderate rain',
        "partly cloudy sky with possibility of moderate rain or thunderstorm",
        "generally cloudy sky with one or two spells of rain or thundershowers",
        "generally cloudy sky with intermittent rain",
        "generally cloudy sky with heavy rain",
        "partly cloudy sky with possibility of heavy rain or thunderstorm",
        "generally cloudy sky with a few spells of rain or thundershowers",
        "generally cloudy sky with continuous rain"
    ]
    snowtext = [
        "snow",
        "rain or snow",
        "light snow",
        "moderate snow",
        "heavy snow"
    ]
    fogmisttext = [
        "shallow fog",
        "fog or mist would occur in early morning",
        "moderate fog",
        "dense fog",
        "very dense fog",
        "fog/mist in the morning and mainly clear sky later",
        "fog or mist in the morning and mainly clear sky later",
        "fog/mist in the morning and partly cloudy sky later",
        "fog or mist in the morning and partly cloudy sky later",
        "haze",
        "mist",
        "heat wave",
        "warn night",
        "heat wave and warm night",
        "cold wave",
        "cold day",
        "cold wave/cold day",
        "dense/very dense fog with cold day conditions",
        "dense or very dense fog with cold day conditions",
        "strong surface winds during day time",
        "chilly winds during day time"
    ]
    iconer = ""
    id = None
    if city in city_Codes.keys():
        id = city_Codes[city]
        page = requests.get("https://mausam.imd.gov.in/imd_latest/contents/LIP/sample4_new.php?id="+str(id))
        soup = BeautifulSoup(page.content, 'html.parser')
        weather=[]
        listo = []
        data = soup.find_all('td')
        for i in range(0,len(data),4):
            t = data[i+3].text.lower()
            if t in cleardaytext:
                iconer = "Clear"
            if t in cloudytext:
                iconer = "Cloudy"
            if t in thunderstormtext:
                iconer = "Thundery"
            if t in raintext:
                iconer = "Rainy"
            if t in snowtext:
                iconer = "Snowy"
            if t in fogmisttext:
                iconer = "Foggy or Misty"
            weather.append([data[i].text,data[i+1].text,data[i+3].text,iconer])
        for i in range(7):
            listo.append(weather[i][1].split("|"))
        return weather, listo
    else:
        error = "No Data Found"
        errortemp = "No Data Found"
        return error, errortemp


@app.route("/searchtab")
def searchtab():
    tableofdata,listofdata=city_weather("Dehradun")
    print(tableofdata)
    if(tableofdata=="No Data Found" and listofdata=="No Data Found"):
        return render_template("405.html"), 405
    else:
        citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
        'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
        'Bhatwari','Yamunotri','Harsil','Gangotri']
        citylistdata = sorted(citylistdata)
        from datetime import date
        today = date.today()
        months = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        d1 = today.strftime("%d")
        d2 = int(today.strftime("%m"))
        dx = ""
        if (d1[1] == "1"):
            d1 = int(d1)
            dx = str(d1) + "st " + months[d2]
        elif (d1[1] == "2"):
            d1 = int(d1)
            dx = str(d1) + "nd " + months[d2]
        elif (d1[1] == "3"):
            d1 = int(d1)
            dx = str(d1) + "rd " + months[d2]
        else:
            d1 = int(d1)
            dx = str(d1) + "th " + months[d2]
        x = 0
        flag = False
        for i in tableofdata:
            if (i[0][:7] == dx[:7]):
                flag = True
                break
            else:
                x = x + 1
        if (flag == False):
            x = 1
        return render_template("index2.html",x = x,lengu = len(tableofdata), citylistdata=citylistdata,city="Dehradun", listofdata=listofdata, tableofdata=tableofdata)



@app.route('/index')
@app.route("/")
def dehradunindex():
    
    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("DEHRADUN")
    print(rainfall)
    print(maxt)
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)
    return render_template("index.html", citylistdata=citylistdata,city="Dehradun", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)


@app.route("/search",methods=['POST'])
def searchbar():
    search_city = request.form['cityname']

    tableofdata, listofdata = city_weather(search_city)

    if (tableofdata == "No Data Found" and listofdata == "No Data Found"):
        return render_template("405.html"), 405
    else:
        citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
        'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
        'Bhatwari','Yamunotri','Harsil','Gangotri']
        citylistdata = sorted(citylistdata)
        from datetime import date
        today = date.today()
        months = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        d1 = today.strftime("%d")
        d2 = int(today.strftime("%m"))
        dx = ""
        if (d1[1] == "1"):
            d1 = int(d1)
            dx = str(d1) + "st " + months[d2]
        elif (d1[1] == "2"):
            d1 = int(d1)
            dx = str(d1) + "nd " + months[d2]
        elif (d1[1] == "3"):
            d1 = int(d1)
            dx = str(d1) + "rd " + months[d2]
        else:
            d1 = int(d1)
            dx = str(d1) + "th " + months[d2]
        x = 0
        flag = False
        for i in tableofdata:
            if (i[0][:7] == dx[:7]):
                flag = True
                break
            else:
                x = x + 1
        if (flag == False):
            x = 1
        return render_template("index2.html",x = x,lengu = len(tableofdata), city=search_city, listofdata=listofdata, tableofdata=tableofdata,citylistdata=citylistdata)



#@app.route("/nowcast")
#def nowcast():
    #import requests
    #from bs4 import BeautifulSoup
    #data = requests.get("http://amssdelhi.gov.in/forecast/NCR/utk_nowcast.html")
    #data_content = data.text
    #soup = BeautifulSoup(data_content, 'html.parser')
    #data_list = soup.find_all('font')
    #dttm = data_list[1].text.split("\n")
    #xttm = data_list[4].text.split("\n")
    #msg = data_list[5].text
    #return render_template("nowcast.html",dttm=dttm,xttm=xttm,msg=msg)

def nowcast():
    data = requests.get("http://amssdelhi.gov.in/forecast/NCR/utk_nowcast.html")
    time.sleep(30)
    data_content = data.text
    soup = BeautifulSoup(data_content, 'html.parser')
    data_list = soup.find_all('font')
    dttm = data_list[1].text.split("\n")
    xttm = data_list[4].text.split("\n")
    msg = data_list[5].text
    import codecs
    with codecs.open('myfile.txt',"w", encoding='utf-8') as f:
        f.write(dttm[0]+"\n"+dttm[1]+"\n"+xttm[0]+" on dated: "+xttm[1]+"\n"+msg+"\n")

@app.route("/haridwar")
def haridwar():

    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("HARIDWAR")
    print(rainfall)
    print(maxt)
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="Haridwar", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)

@app.route("/dehradun")
def dehradun():
    
    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("DEHRADUN")
    print(rainfall)
    print(maxt)
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="Dehradun", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/mussoorie")
def mussoorie():
    tableofdata, listofdata = city_weather("Mussoorie")

    if (tableofdata == "No Data Found" and listofdata == "No Data Found"):
        return render_template("405.html"), 405
    else:
        citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
        'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
        'Bhatwari','Yamunotri','Harsil','Gangotri']
        citylistdata = sorted(citylistdata)

        return render_template("index.html",citylistdata=citylistdata, city="Mussoorie", listofdata=listofdata, tableofdata=tableofdata)


@app.route("/tehri")
def tehri():

    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("TEHRI-GARHWAL")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="TEHRI_GARHWAL", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)


@app.route("/pauri")
def pauri():



    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("PAURI")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="PAURI", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)


@app.route("/almora")
def almora():

    
    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("ALMORA")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="ALMORA", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)


@app.route("/mukteshwar")
def mukteshwar():

    tableofdata,listofdata=city_weather("Mukteshwar")
    if(tableofdata=="No Data Found" and listofdata=="No Data Found"):
        return render_template("405.html"), 405
    else:
        citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
        'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
        'Bhatwari','Yamunotri','Harsil','Gangotri']
        citylistdata = sorted(citylistdata)

        return render_template("index.html", citylistdata=citylistdata,city="Mukteshwar",  listofdata=listofdata, tableofdata=tableofdata)


@app.route("/nainital")
def nainital():


    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("NAINITAL")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="NAINITAL", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/udhamsinghnagar")
def udhamsinghnagar():



    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("UDHAMSINGH-NGR")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="UDHAMSINGH-NGR", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/champawat")
def champawat():



    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("CHAMPAWAT")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="CHAMPAWAT", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/pithoragarh")
def pithoragarh():

    

    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("PITHORAGARH")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="PITHORAGARH", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/chamoli")
def chamoli():


    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("CHAMOLI")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="CHAMOLI", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)



@app.route("/joshimath")
def joshimath():
    tableofdata, listofdata = city_weather("Joshimath")

    if (tableofdata == "No Data Found" and listofdata == "No Data Found"):
        return render_template("405.html"), 405
    else:
        citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
        'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
        'Bhatwari','Yamunotri','Harsil','Gangotri']
        citylistdata = sorted(citylistdata)

        return render_template("index.html", citylistdata=citylistdata,city="Joshimath", listofdata=listofdata, tableofdata=tableofdata)



@app.route("/uttarkashi")
def uttarkashi():


    rainfall, maxt, mint, tcc, marh, mirh, ws, wd = show_Forcast("UTTARKASHI")
    citylistdata = ['Pithoragarh','Pantnagar','Nainital','Hemkund Sahib','Champawat','Mukteshwar','Almora','Chamoli','Joshimath','Govindghat','Badrinath','Devprayag','Srinagar','Rudraprayag',
    'Gauchar','Augustmuni','Ukhimat','Soneprayag','Kedarnath','Haridwar','Rishikesh','Chamba','Tehri','Dehradun','Mussoorie','Nainbagh','Chinyalisaur','Lakhamandal','Barkot','Uttarkashi',
    'Bhatwari','Yamunotri','Harsil','Gangotri']
    citylistdata = sorted(citylistdata)

    return render_template("district.html", citylistdata=citylistdata,city="UTTARKASHI", rainfall=rainfall, maxt=maxt, mint=mint, tcc=tcc, marh=marh, mirh=mirh, ws=ws, wd=wd)


@app.errorhandler(404)
def error404(error):
    return render_template("404.html"), 404

@app.errorhandler(500)
def error500(error):
    return render_template("500.html"), 500

@app.errorhandler(405)
def error405(error):
    return render_template("405.html"), 405

#@app.route('/news')
#def news():
    #news, desc, date = [], [], []
    #res = requests.get("https://timesofindia.indiatimes.com/topic/monsoon")
    #soup = BeautifulSoup(res.text,'html.parser')
    #d = soup.find('ul', {'itemprop': 'ItemList'})
    #for i in d.findAll('li', {'class': 'article'}):
        #k = i.find('span', {'class': 'title'})
        #l = i.find('p')
        #m = i.find('span', {'class': 'meta'})
        #news.append(k.text)
        #desc.append(l.text)
        #date.append(m.text[0:10])
    #return render_template('news.html', news=news, desc=desc, date=date)

@app.route("/contact")
def contact():
    return render_template("contact.html")

def isParenthesis(c): 
    return ((c == '(') or (c == ')'))
def isValidString(str): 
    cnt = 0
    for i in range(len(str)): 
        if (str[i] == '('): 
            cnt += 1
        elif (str[i] == ')'): 
            cnt -= 1
        if (cnt < 0): 
            return False
    return (cnt == 0) 
def removeInvalidParenthesis(str): 
    if (len(str) == 0): 
        return 
    visit = set() 
    q = [] 
    temp = 0
    level = 0
    q.append(str) 
    visit.add(str) 
    while(len(q)): 
        str = q[0] 
        q.pop() 
        if (isValidString(str)):  
            level = True
        if (level): 
            continue
        for i in range(len(str)): 
            if (not isParenthesis(str[i])): 
                continue
            temp = str[0:i] + str[i + 1:] 
            if temp not in visit: 
                q.append(temp) 
                visit.add(temp) 
                return temp

def alertsandwarnings():
    import PyPDF2
    pdfFileObj = open(r'DISTRICT_WARNING.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj, strict=False)
    number_of_pages =pdfReader.getNumPages()
    pageObj = pdfReader.getPage(0)
    page_content = pageObj.extractText()
    watch = "WATCH: "
    alert = ""
    f = 1
    page_data1 = page_content.split("\n")
    watch = ""
    flagw = False
    alert = ""
    flaga = False
    taction = ""
    flagta = False
    removlist = ["' '"," ","''","'",'']
    for i in page_data1:
        if i in removlist:
            page_data1.remove(i)
    for i in range(len(page_data1)):
        if("WATCH" in page_data1[i] and len(page_data1[i])<=10):
            flagw=True
            continue
        if(flagw==True):
            watch = watch + page_data1[i]
            if(len(page_data1[i])>0 and page_data1[i][-1]=="."):
                break;

    for i in range(len(page_data1)):
        if("ALERT" in page_data1[i] and len(page_data1[i])<=10):
            flaga=True
            continue
        if(flaga==True):
            alert = alert + page_data1[i]
            if(len(page_data1[i])>0 and page_data1[i][-1]=="."):
                break;
    for i in range(len(page_data1)):
        if("TAKE ACTION" in page_data1[i] and len(page_data1[i])<=20):
            flagta=True
            continue
        if(flagta==True):
            taction = taction + page_data1[i]
            if(len(page_data1[i])>0 and page_data1[i][-1]=="."):
                break;
                
    newwatch=removeInvalidParenthesis(watch)
    newalter=removeInvalidParenthesis(alert)
    newtakea=removeInvalidParenthesis(taction)
    if(newwatch==None):
        newwatch = watch
    if(newalter==None):
        newalter = alert
    if(newtakea==None):
        newtakea = taction

    return newwatch,newalter,newtakea

def forcast_warn_funct(cityfind):
    from docx.api import Document
    document = Document('DISTRICT_WARNING.docx')
    #Logic for Overall Warning

    table1 = document.tables[2]
    data = []
    keys = None
    for i, row in enumerate(table1.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

    Warning = list(data[0].values())[0].split('\n')

    for i in Warning:
        print(i)
    #Logic for District-wise Warning
    table = document.tables[0]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        rows_data = dict(zip(keys, text))
        data.append(rows_data)


    districtWise = []

    for i in range(13):
        tempData = list(data[i].values())
        if(tempData[0]==cityfind):
            districtWise = tempData[1:]
            break
    return Warning,districtWise

@app.route("/forecast-warning")
def forecast_warning():
    citylistdata = ["Almora","Bageshwar","Chamoli","Champawat","Dehradun","Haridwar","Nainital","Pauri Garhwal","Pithoragarh","Rudraprayag","Tehri Garhwal","Udham Singh Nagar","Uttarkashi"]
    citylistdata = sorted(citylistdata)
    #w,a,t = alertsandwarnings()
    warn,distw = forcast_warn_funct("Dehradun")
    #print("tatatatata : ",t)
    return render_template("districtlevelfor.html", city='Dehradun',warn=warn, distw=distw,citylistdata=citylistdata)

@app.route('/fwsearch',methods=['POST'])
def fwsearch():
    search_city = request.form['cityname']
    citylistdata = ["Almora","Bageshwar","Chamoli","Champawat","Dehradun","Haridwar","Nainital","Pauri Garhwal","Pithoragarh","Rudraprayag","Tehri Garhwal","Udham Singh Nagar","Uttarkashi"]
    citylistdata = sorted(citylistdata)
    warn,distw = forcast_warn_funct(search_city)
    return render_template("districtlevelfor.html", city=search_city,warn=warn, distw=distw,citylistdata=citylistdata)

@app.route("/fw")
def fw():
    #pdfdownloader()
    alert,watch = alertsandwarnings()
    return render_template("for.html", alert=alert, watch=watch)

@app.route("/video")
def video():
    from firebase import firebase
    firebase = firebase.FirebaseApplication('https://mausamkendradehradun.firebaseio.com/', None)
    results = firebase.get('mausamkendradehradun/videolibrary', '')
    video_url = []
    time = []
    thumbnail = []
    date = []
    for i in results.values():
        video_url.append(i['video_url'])
        thumbnail.append(i['thumbnail'])
        date.append(i['date'])
        time.append(i['time'])
    return render_template("video.html",leng = len(date) ,thumbnail=thumbnail, video_url=video_url, datee = date, timee=time)



# ==============================================ADMIN SECTION=========================================================




#@app.route("/admin_upload_video_forecast")
#def video12():
    #return render_template("admin_upload_video_forercast.html")

#@app.route("/upload_video", methods = ["GET", "POST"])
#def upload_video():
    #if request.method == "POST":
        #f = request.files['file']
        #f2 = request.files['file2']
        #if(f.filename.endswith('.mp4')):
            #from datetime import date, datetime
            #today = date.today()
            #now = datetime.now()
            #date = today.strftime("%d-%m-%Y")
            #time = now.strftime("%H:%M:%S")
            #time22 = now.strftime("%H-%M-%S")
            #fname = '/home/ubuntu/flaskproject/static/video/' + 'si_bkrm_jgt_mcdun_btin_' + date + "_" + time22  + ".mp4"
            #fname2 = '/home/ubuntu/flaskproject/static/thumb/' + 'si_bkrm_jgt_mcdun_btin_' + date   + "_" + time22  + ".jpg"
            #f.save(fname)
            #f2.save(fname2)
            #video_url = "/home/ubuntu/flaskproject/static/video/" + 'si_bkrm_jgt_mcdun_btin_' + date + "_" + time22  + ".mp4"
            #thumbnail = "/home/ubuntu/flaskproject/static/thumb/" + 'si_bkrm_jgt_mcdun_btin_' + date +"_" + time22   + ".jpg"
            #sqliteConnection = sql.connect('database.db')
            #cursor = sqliteConnection.cursor()
            #cursor.execute("INSERT INTO video (srno,video_url,thumbnail,date,time) VALUES(?, ?, ?, ?, ?)",
            #               (None, video_url, thumbnail, date, time))
            #sqliteConnection.commit()
            #cursor.close()
            #cur = mysql.connection.cursor()
            #cur.execute("INSERT INTO video (srno,video_url,thumbnail,date,time) VALUES(%s, %s, %s, %s, %s)",
            #               (None, video_url, thumbnail, date, time))
            #mysql.connection.commit()
            #cur.close()
            #return render_template("admin_upload_video_forercast.html", msg="success")
        #else:
            #return render_template("admin_upload_video_forercast.html", msg="failed_video")
    #else:
        #return render_template("admin_upload_video_forercast.html", msg="failed")

#@app.route("/adminuploaddatawarning")
#def adminuploaddatawarning():
    #return render_template("adminuploaddatawarning.html", msg="firsttime")

@app.route("/nowcast")
def nowcstr():
    import codecs
    with codecs.open('myfile.txt', "r", encoding='utf-8') as f:
        rs = f.read()
    ms = rs.split("\n")
    return render_template("nowcast.html",ms=ms)

@app.route("/pdfdownloader",methods=["GET","POST"])
def pdfdownloader():
    if request.method=="POST":
        f = request.files['file']
        if (f.filename.endswith('.pdf') or f.filename.endswith('.docx')):
            fname=f.filename
            f.save(fname)
            return render_template("adminuploaddatawarning.html", msg="success")
        else:
            return render_template("adminuploaddatawarning.html", msg="failed_pdf")
    else:
        return render_template("adminuploaddatawarning.html", msg="failed")

@app.route("/admin_upload_video_forecast")
def video12():
    if 'username' in session:
        return render_template("admin_upload_video_forercast.html")
    else:
        return render_template('login.html')

@app.route("/upload_video", methods = ["GET", "POST"])
def upload_video():
    if 'username' in session:
        if request.method == "POST":
            f = request.files['file']
            f2 = request.files['file2']
            if(f.filename.endswith('.mp4')):
                from datetime import date, datetime
                from firebase import firebase
                firebase = firebase.FirebaseApplication('https://mausamkendradehradun.firebaseio.com/',None)
                today = date.today()
                now = datetime.now()
                date = today.strftime("%d-%m-%Y")
                time = now.strftime("%H:%M:%S")
                time22 = now.strftime("%H-%M-%S")
                fname = '/home/ubuntu/flaskproject/static/video/' + 'si_bkrm_jgt_mcdun_btin_' + date + "_" + time22  + ".mp4"
                fname2 = '/home/ubuntu/flaskproject/static/thumb/' + 'si_bkrm_jgt_mcdun_btin_' + date   + "_" + time22  + ".jpg"
                f.save(fname)
                f2.save(fname2)
                video_url = "/home/ubuntu/flaskproject/static/video/" + 'si_bkrm_jgt_mcdun_btin_' + date + "_" + time22  + ".mp4"
                thumbnail = "/home/ubuntu/flaskproject/static/thumb/" + 'si_bkrm_jgt_mcdun_btin_' + date +"_" + time22   + ".jpg"
                data={'video_url': video_url,'thumbnail': thumbnail,'date': date,'time': time}
                results = firebase.get('mausamkendradehradun/videolibrary','')
                ant = []
                try:
                    for i in results.values():
                        ant.append(i['video_url'])
                    if video_url not in ant:
                        resulti = firebase.post('mausamkendradehradun/videolibrary', data)
                        print("DATA INSERTED")
                    else:
                        print("ALREADY INSERTED")
                except AttributeError:
                    resulti = firebase.post('mausamkendradehradun/videolibrary', data)
                    print("DATA INSERTED")
                return render_template("admin_upload_video_forercast.html", msg="success")
            else:
                return render_template("admin_upload_video_forercast.html", msg="failed_video")
        else:
            return render_template("admin_upload_video_forercast.html", msg="failed")
    else:
        return render_template('login.html')


@app.route("/adminuploaddatawarning")
def adminuploaddatawarning():
    if 'username' in session:
        return render_template("adminuploaddatawarning.html", msg="firsttime")
    else:
        return render_template('login.html')

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/adminloginmausamvibhag', methods=['GET', 'POST'])
def loginmausam():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if(username=="bikramjagatadminlog@mausam.com" and password == "bikramjagatadminlog@99"):
            session['username'] = username
            return redirect(url_for('adminuploaddatawarning'))
    return render_template('login.html',msg = "UsernamePassword")

@app.route('/logout')
def logout():
   # remove the username from the session if it is there
   session.pop('username', None)
   return redirect(url_for('login'))
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within parent,
    in document order. Each returned value is an instance of either Table or
    Paragraph. parent would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


scheduler = BackgroundScheduler()
scheduler.add_job(func=nowcast, trigger="interval",hours=1)
scheduler.start()
atexit.register(lambda: scheduler.shutdown())

if __name__ == '__main__':
    app.run(host="127.0.0.1",port=5000)
    
